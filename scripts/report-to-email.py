#!/usr/bin/env python3
"""
report-to-email.py — genera PDF + DOCX de un reporte y lo manda por email.
Reusable para TODO reporte/propuesta/informe de Wenu Mapu.

  echo "texto..." | python3 scripts/report-to-email.py "Título del reporte"
  python3 scripts/report-to-email.py "Título" archivo.md
  ... [destinatario opcional]  (default wenu.mapu8@gmail.com)

Lee SMTP de ~/.hermes/.env (SMTP_HOST/PORT/USER/PASS). Read-only salvo el envío.
"""
import sys, os, re, smtplib, ssl, datetime
from email.message import EmailMessage
from pathlib import Path

HOME = str(Path.home())
DEFAULT_TO = "wenu.mapu8@gmail.com"

def load_env(p):
    o = {}
    try:
        for line in open(p, encoding="utf-8"):
            m = re.match(r"^\s*([A-Z_][A-Z0-9_]*)\s*=\s*(.*)\s*$", line)
            if m:
                v = m.group(2).strip()
                if len(v) >= 2 and v[0] in "\"'" and v[-1] == v[0]:
                    v = v[1:-1]
                o[m.group(1)] = v
    except Exception:
        pass
    return o

def make_pdf(path, title, body):
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    styles = getSampleStyleSheet()
    h = ParagraphStyle("h", parent=styles["Title"], fontSize=18, textColor=colors.HexColor("#6a4a28"), spaceAfter=6)
    sub = ParagraphStyle("sub", parent=styles["Normal"], fontSize=9, textColor=colors.grey, spaceAfter=14)
    hd = ParagraphStyle("hd", parent=styles["Heading2"], fontSize=13, textColor=colors.HexColor("#0a0a0a"), spaceBefore=10, spaceAfter=4)
    body_s = ParagraphStyle("b", parent=styles["Normal"], fontSize=10.5, leading=15)
    doc = SimpleDocTemplate(path, pagesize=LETTER, topMargin=0.8*inch, bottomMargin=0.8*inch, leftMargin=0.9*inch, rightMargin=0.9*inch)
    flow = [Paragraph(title, h), Paragraph("Wenu Mapu · " + datetime.date.today().isoformat(), sub)]
    for raw in body.split("\n"):
        line = raw.rstrip()
        if not line:
            flow.append(Spacer(1, 5)); continue
        esc = (line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
        esc = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", esc)
        if re.match(r"^#{1,6}\s", line):
            flow.append(Paragraph(re.sub(r"^#{1,6}\s", "", esc), hd))
        elif re.match(r"^\s*[-*•]\s", line):
            flow.append(Paragraph("• " + re.sub(r"^\s*[-*•]\s", "", esc), body_s))
        else:
            flow.append(Paragraph(esc, body_s))
    doc.build(flow)

def make_docx(path, title, body):
    from docx import Document
    from docx.shared import Pt, RGBColor
    d = Document()
    t = d.add_heading(title, level=0)
    d.add_paragraph("Wenu Mapu · " + datetime.date.today().isoformat()).italic = True
    for raw in body.split("\n"):
        line = raw.rstrip()
        if not line:
            continue
        if re.match(r"^#{1,6}\s", line):
            d.add_heading(re.sub(r"^#{1,6}\s", "", line), level=2)
        elif re.match(r"^\s*[-*•]\s", line):
            d.add_paragraph(re.sub(r"^\s*[-*•]\s", "", line), style="List Bullet")
        else:
            p = d.add_paragraph()
            for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", line)):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True
    d.save(path)

def main():
    if len(sys.argv) < 2:
        print("uso: report-to-email.py <título> [archivo|-] [destinatario]"); sys.exit(1)
    title = sys.argv[1]
    src = sys.argv[2] if len(sys.argv) > 2 else "-"
    to = sys.argv[3] if len(sys.argv) > 3 else DEFAULT_TO
    body = sys.stdin.read() if src == "-" else open(src, encoding="utf-8").read()

    slug = re.sub(r"[^a-z0-9]+", "-", title.lower()).strip("-")[:50] or "reporte"
    stamp = datetime.date.today().isoformat()
    outdir = HOME + "/wenu-frontend/reports"
    os.makedirs(outdir, exist_ok=True)
    pdf = f"{outdir}/{stamp}-{slug}.pdf"
    docx = f"{outdir}/{stamp}-{slug}.docx"
    make_pdf(pdf, title, body); make_docx(docx, title, body)
    print(f"generado: {pdf} + {docx}")

    env = load_env(HOME + "/.hermes/.env")
    host, port = env.get("SMTP_HOST"), int(env.get("SMTP_PORT", "587"))
    user, pw = env.get("SMTP_USER"), env.get("SMTP_PASS")
    if not (host and user and pw):
        print("SMTP incompleto — PDF/DOCX quedaron en disco, no se envió."); sys.exit(2)
    msg = EmailMessage()
    msg["Subject"] = f"[Wenu Mapu] {title} — {stamp}"
    msg["From"] = user
    msg["To"] = to
    msg.set_content(f"Adjunto el reporte '{title}' en PDF y DOCX.\n\nGenerado automáticamente — Wenu Mapu / Claude.\n\n{body[:1500]}")
    for f in (pdf, docx):
        with open(f, "rb") as fh:
            data = fh.read()
        sub = "pdf" if f.endswith("pdf") else "vnd.openxmlformats-officedocument.wordprocessingml.document"
        maintype = "application" if f.endswith("pdf") else "application"
        msg.add_attachment(data, maintype=maintype, subtype=sub, filename=os.path.basename(f))
    ctx = ssl.create_default_context()
    with smtplib.SMTP(host, port, timeout=30) as s:
        s.starttls(context=ctx)
        s.login(user, pw)
        s.send_message(msg)
    print(f"email enviado a {to} (PDF+DOCX) vía {host}")

if __name__ == "__main__":
    main()
